from typing import Dict, Any
from docx import Document  # type: ignore

def parse(filepath: str) -> Dict[str, Any]:
    try:
        # Load the DOCX document
        doc = Document(filepath)  # type: ignore
        content = ""
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:  # type: ignore
            paragraph_text = paragraph.text  # type: ignore
            if paragraph_text and paragraph_text.strip():
                content += paragraph_text + "\n"
        
        # Extract text from tables
        for table in doc.tables:  # type: ignore
            for row in table.rows:  # type: ignore
                for cell in row.cells:  # type: ignore
                    cell_text = cell.text  # type: ignore
                    if cell_text and cell_text.strip():
                        content += cell_text + " "
                content += "\n"
        
        return {
            "status": "success",
            "filename": filepath,
            "content": content.strip(),
            "error": None
        }
    except Exception as e:
        return {
            "status": "error",
            "filename": filepath,
            "content": None,
            "error": str(e)
        }